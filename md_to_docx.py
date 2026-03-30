#!/usr/bin/env python3
"""
Simple Markdown to DOCX converter focused on legal/professional docs.

Usage:
    python md_to_docx.py --input markdown/marco_legal_formateado.md --output output/marco_legal_formateado.docx
"""

from __future__ import annotations

import argparse
import os
import re
from pathlib import Path

from docx import Document


HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
BULLET_RE = re.compile(r"^\s*[-*]\s+(.*)$")
NUMBER_RE = re.compile(r"^\s*\d+[\.)]\s+(.*)$")


def load_dotenv(dotenv_path: Path = Path(".env")) -> None:
    """Load KEY=VALUE pairs from .env into process env if not already set."""
    if not dotenv_path.exists():
        return

    for raw_line in dotenv_path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        key = key.strip()
        value = value.strip().strip('"').strip("'")
        if key and key not in os.environ:
            os.environ[key] = value


def convert_md_to_docx(input_path: Path, output_path: Path) -> None:
    text = input_path.read_text(encoding="utf-8")
    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")

    doc = Document()

    for raw_line in lines:
        line = raw_line.rstrip()

        if not line.strip():
            doc.add_paragraph("")
            continue

        if line.strip() == "---":
            # Skip markdown horizontal rules
            continue

        heading_match = HEADING_RE.match(line)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            title = heading_match.group(2).strip()
            doc.add_heading(title, level=level)
            continue

        bullet_match = BULLET_RE.match(line)
        if bullet_match:
            doc.add_paragraph(bullet_match.group(1).strip(), style="List Bullet")
            continue

        number_match = NUMBER_RE.match(line)
        if number_match:
            doc.add_paragraph(number_match.group(1).strip(), style="List Number")
            continue

        # Plain paragraph
        doc.add_paragraph(line)

    doc.save(str(output_path))


def main() -> None:
    load_dotenv()

    default_input = os.getenv("DOCX_INPUT_PATH", "markdown/marco_legal_formateado.md")
    default_output = os.getenv("DOCX_OUTPUT_PATH", "output/marco_legal_formateado.docx")

    parser = argparse.ArgumentParser(description="Convert Markdown file to DOCX")
    parser.add_argument(
        "--input",
        "-i",
        default=default_input,
        help="Input .md file",
    )
    parser.add_argument(
        "--output",
        "-o",
        default=default_output,
        help="Output .docx file",
    )
    args = parser.parse_args()

    input_path = Path(args.input)
    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if not input_path.exists():
        raise SystemExit(f"Input file not found: {input_path}")

    if input_path.suffix.lower() != ".md":
        raise SystemExit("Input must be a .md file")

    if output_path.suffix.lower() != ".docx":
        raise SystemExit("Output must be a .docx file")

    convert_md_to_docx(input_path, output_path)
    print(f"OK: Generated {output_path}")


if __name__ == "__main__":
    main()
